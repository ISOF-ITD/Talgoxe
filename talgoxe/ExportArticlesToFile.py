import datetime
import docx
import ezodf
from django.conf import settings
from os import chdir, popen, rename, environ
from os.path import abspath, dirname, join
from shutil import copyfile
from talgoxe.ArticleManager import ArticleManager
from talgoxe.common_functions import *
from tempfile import mkdtemp


class ExportArticlesToFile:
    def __init__(self, format):
        self.format = format
        initialisers = {
            'pdf' : self.start_pdf,
            'odt' : self.start_odf,
            'docx' : self.start_docx,
        }

        generators = {
            'pdf' : self.generate_pdf_paragraph,
            'odt' : self.generate_odf_paragraph,
            'docx' : self.generate_docx_paragraph,
        }

        savers = {
            'pdf' : self.save_pdf,
            'odt' : self.save_odf,
            'docx' : self.save_docx,
        }

        self.start_document = initialisers[format]
        self.generate_paragraph = generators[format]
        self.save_document = savers[format]

    def export_letter(self, letter):
        # Get data from database.
        articles = ArticleManager.get_articles_by_letter(letter)
        if len(articles) >= 1:
            # Create file with article information.
            self.dirname = mkdtemp('', 'SDLartikel')
            self.filename = join(self.dirname, 'sdl.%s' % self.format)
            self.start_document()
            filename = letter + datetime.datetime.now().strftime("%Y%m%d") + '.' + self.format
            for article in articles:
                self.generate_paragraph(article)
            self.save_document()
            staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
            rename(self.filename, join(staticpath, filename))

    # Create files, one file for each letter in the swedish alphabet.
    def export_letters(self):
        alphabet = get_swedish_alphabet()
        for character in alphabet:
            self.export_letter(character)

    def generate_pdf_paragraph(self, artikel):
        self.document.write("\\hskip-0.5em")
        if artikel.rang > 0:
            self.document.write('\lohi[left]{}{\SDL:SO{%d}}' % artikel.rang) # TODO Real superscript
        self.document.write("\\SDL:SO{%s}" % artikel.lemma)
        for segment in artikel.fjädrar: # TODO Handle moments!  segment.ismoment and segment.display
            if not segment.preventspace and not segment.isrightdelim():
                self.document.write(' ') # FIXME But not if previous segment is left delim!
            type = segment.typ
            text = segment.format().replace(u'\\', '\\textbackslash ').replace('~', '{\\char"7E}')
            self.document.write(('\\SDL:%s{%s}' % (type, text)))
        self.document.write("\\par")

    def start_pdf(self):
        self.document = open(self.filename.replace('.pdf', '.tex'), 'w')
        self.document.write("\\mainlanguage[sv]")
        self.document.write("\\setupbodyfont[pagella, 12pt]\n")
        self.document.write("\\setuppagenumbering[state=stop]\n")
        with open(join(dirname(abspath(__file__)), 'lib', 'sdl-setup.tex'), 'r', encoding = 'UTF-8') as file:
            self.document.write(file.read())

        self.document.write("""
            \\starttext
            """)

        self.document.write("\\startcolumns[n=2,balance=yes]\n")

    def save_pdf(self):
        self.document.write("\\stopcolumns\n")
        self.document.write("\\stoptext\n")
        self.document.close()

        chdir(self.dirname)
        environ['PATH'] = "%s:%s" % (settings.CONTEXT_PATH, environ['PATH'])
        environ['TMPDIR'] = '/tmp'
        popen("context --batchmode sdl.tex").read()

    def add_odf_style(self, type, xmlchunk):
        self.document.inject_style("""
            <style:style style:name="%s" style:family="text">
                <style:text-properties %s />
            </style:style>
            """ % (type, xmlchunk))

    def start_odf(self):
        self.document = ezodf.newdoc(doctype = 'odt', filename = self.filename)
        self.add_odf_style('SO', 'fo:font-weight="bold"')
        self.add_odf_style('OK', 'fo:font-size="9pt"')
        self.add_odf_style('G', 'fo:font-size="9pt"')
        self.add_odf_style('DSP', 'fo:font-style="italic"')
        self.add_odf_style('TIP', 'fo:font-size="9pt"')
        # IP
        self.add_odf_style('M1', 'fo:font-weight="bold"')
        self.add_odf_style('M2', 'fo:font-weight="bold"')
        # VH, HH, VR, HR
        self.add_odf_style('REF', 'fo:font-weight="bold"')
        self.add_odf_style('FO', 'fo:font-style="italic"')
        self.add_odf_style('TIK', 'fo:font-style="italic" fo:font-size="9pt"')
        self.add_odf_style('FLV', 'fo:font-weight="bold" fo:font-size="9pt"')
        # ÖVP. Se nedan!
        # BE, ÖV
        # ÄV, ÄVK se nedan
        # FOT
        self.add_odf_style('GT', 'fo:font-size="9pt"')
        self.add_odf_style('SOV', 'fo:font-weight="bold"')
        # TI
        # HV, INT
        self.add_odf_style('OKT', 'fo:font-size="9pt"')
        # VS
        self.add_odf_style('GÖ', 'fo:font-size="9pt"')
        # GP, UST
        self.add_odf_style('US', 'fo:font-style="italic"')
        # GÖP, GTP, NYR, VB
        self.add_odf_style('OG', 'style:text-line-through-style="solid"')
        self.add_odf_style('SP', 'fo:font-style="italic"')

    def save_odf(self):
        self.document.save()

    def generate_odf_paragraph(self, artikel):
        paragraph = ezodf.Paragraph()
        paragraph += ezodf.Span(artikel.lemma, style_name = 'SO') # TODO Homografnumrering!
        for segment in artikel.fjädrar:
            type = segment.typ
            if not type == 'KO':
                if not segment.preventspace and not segment.isrightdelim():
                    paragraph += ezodf.Span(' ')
                paragraph += ezodf.Span(segment.format(), style_name = type)
        self.document.body += paragraph

    def start_docx(self):
        self.document = docx.Document()
        self.add_docx_styles()
        return self.document

    def save_docx(self):
        self.document.save(self.filename)

    def generate_docx_paragraph(self, artikel):
        paragraph = self.document.add_paragraph()
        if artikel.rang > 0:
            paragraph.add_run(str(artikel.rang), style = 'SO').font.superscript = True
        paragraph.add_run(artikel.lemma, style = 'SO')
        for segment in artikel.fjädrar:
            type = segment.typ
            if not type == 'KO':
                if not segment.preventspace and not segment.isrightdelim():
                    paragraph.add_run(' ', style = self.document.styles[type])
                if type == 'ÖVP': # TODO Något bättre
                    run = '(' + segment.format() + ')'
                else:
                    run = segment.format()
                paragraph.add_run(run, style = self.document.styles[type])

    def add_docx_styles(self): # TODO Keyword arguments?
        self.add_docx_style('SO', False, True, 12)
        self.add_docx_style('OK', False, False, 9)
        self.add_docx_style('G', False, False, 9)
        self.add_docx_style('DSP', True, False, 12)
        self.add_docx_style('TIP', False, False, 9)
        self.add_docx_style('IP', False, False, 12)
        self.add_docx_style('M1', False, True, 12)
        self.add_docx_style('M2', False, True, 12)
        self.add_docx_style('VH', False, False, 12)
        self.add_docx_style('HH', False, False, 12)
        self.add_docx_style('VR', False, False, 12)
        self.add_docx_style('HR', False, False, 12)
        self.add_docx_style('REF', False, True, 12)
        self.add_docx_style('FO', True, False, 12)
        self.add_docx_style('TIK', True, False, 9)
        self.add_docx_style('FLV', False, True, 9)
        self.add_docx_style('ÖVP', False, False, 12)
        self.add_docx_style('BE', False, False, 12)
        self.add_docx_style('ÖV', False, False, 12)
        self.add_docx_style('ÄV', False, False, 12) # FIXME Skriv “även” :-)
        self.add_docx_style('ÄVK', True, False, 12) # FIXME Skriv även även här ;-)
        self.add_docx_style('FOT', True, False, 12)
        self.add_docx_style('GT', False, False, 9)
        self.add_docx_style('SOV', False, True, 12)
        for style in ('TI', 'HV', 'INT'):
            self.add_docx_style(style)
        self.add_docx_style('OKT', False, False, 9)
        self.add_docx_style('VS')
        self.add_docx_style('GÖ', False, False, 9)
        self.add_docx_style('GP')
        self.add_docx_style('UST', True, False, 12)
        self.add_docx_style('US', True, False, 12)
        for style in ('GÖP', 'GTP', 'NYR', 'VB'):
            self.add_docx_style(style)
        OG = self.document.styles.add_style('OG', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        OG.font.strike = True
        self.add_docx_style('SP', True, False, 12)
        self.add_docx_style('M0', False, True, 18)
        self.add_docx_style('M3', True, False, 6)

    def add_docx_style(self, type, italic = False, bold = False, size = 12):
        style = self.document.styles.add_style(type, docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        style.base_style = self.document.styles['Default Paragraph Font']
        if italic:
            style.font.italic = True
        if bold:
            style.font.bold = True
        style.font.size = docx.shared.Pt(size)

    def export(self, article_ids):
        self.dirname = mkdtemp('', 'SDLartikel')
        self.filename = join(self.dirname, 'sdl.%s' % self.format)
        self.start_document()
        if (len(article_ids) == 0) or (len(article_ids) == 1 and article_ids[0] == ''):
            # Nothing to do. No articles should be output to file.
            pass
        elif len(article_ids) == 1:
            article = ArticleManager.get_article(article_ids[0])
            self.generate_paragraph(article)
            filename = '%s-%s.%s' % (article_ids[0], article.lemma, self.format)
        elif len(article_ids) > 1:
            # Write article information to file.
            articles = ArticleManager.get_articles_by_ids(article_ids)
            filename = 'sdl-utdrag.%s' % self.format # FIXME Unikt namn osv.
            for article in articles:
                self.generate_paragraph(article)
        self.save_document()
        staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
        copyfile(self.filename, join(staticpath, filename))

        return join('ord', filename)

    def export_articles(self, articles, user_name):
        # Get temp file.
        self.dirname = mkdtemp('', 'SDLartikel')
        self.filename = join(self.dirname, 'sdl.%s' % self.format)
        self.start_document()

        # Write article information to file.
        filename = f'{user_name}-sdl.{self.format}'
        for article in articles:
            self.generate_paragraph(article)
        self.save_document()

        # Copy temp file to output file.
        staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
        filepath = join(staticpath, filename)
        copyfile(self.filename, filepath)

        return filepath

