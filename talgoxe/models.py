# -*- coding: utf-8 -*-
from __future__ import unicode_literals
import datetime
from math import floor
from os import chdir, popen, rename, environ
from os.path import abspath, dirname, join
from re import match, split
from tempfile import mkdtemp
from django.contrib.auth.models import User
from django.db.models import Max

import ezodf
import docx
from shutil import copyfile

from django.conf import settings
from django.core.exceptions import ObjectDoesNotExist, PermissionDenied
from django.db import models

# FIXME Superscripts!

def is_empty_string(string):
    if (string is None):
        return True
    return (len(string) < 1)

def is_not_empty_string(string):
    return not is_empty_string(string)

class UnsupportedFormat(Exception):
  pass

edit_users = {}

class AccessManager:

    @staticmethod
    def check_edit_permission(user):
        AccessManager.init_access_control()

        # Check if user has permission to edit articles.
        if not (user.username in edit_users):
            # User is not in group 'TalgoxenEdit'.
            raise PermissionDenied(f"User {user.first_name} {user.last_name} ({user.username}) does not have edit permission.")

    @staticmethod
    def init_access_control():
        if len(edit_users) == 0:
            # Get all users with edit permission.
            editors = User.objects.filter(groups__name='TalgoxenEdit')
            for editor in editors:
                edit_users[editor.username] = editor

class Artikel(models.Model):
    class Meta:
        ordering = ('lemma_sortable', 'rang')

    lemma = models.CharField(max_length = 100)
    lemma_sortable = models.CharField(max_length = 100)
    rang = models.SmallIntegerField()
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)
    segments = []

    superscripts = [
        '⁰', '¹', '²', '³', '⁴', '⁵', '⁶', '⁷', '⁸', '⁹',
    ]

    def __str__(self):
        r = self.rang
        prefix = ''
        while r > 0:
            prefix = Artikel.superscripts[r % 10] + prefix
            r = floor(r / 10)
        return prefix + self.lemma

    @staticmethod
    def create(post_data):
        nylemma = post_data['stickord']
        företrädare = Artikel.objects.filter(lemma = nylemma)
        maxrang = företrädare.aggregate(Max('rang'))['rang__max']
        if maxrang == None:
            rang = 0
        elif maxrang == 0:
            artikel0 = företrädare.first()
            artikel0.rang = 1
            artikel0.save()
            rang = 2
        elif maxrang > 0:
            rang = maxrang + 1
        article = Artikel.objects.create(lemma = nylemma, lemma_sortable = Artikel.get_lemma_sortable(nylemma), rang = rang)
        article.update(post_data)
        return article

    def get_spole(self, i):
        if i < len(self.spolar()):
            return self.spolar()[i] # Man kan missa vissa värden av pos, så .get(pos = i) funkar inte. Se t.ex. 1öden (id 3012683), spole saknas för pos = 2. AR 2018-06-03.
        else:
            return None

    @staticmethod
    def get_lemma_sortable(lemma):
        lemma_sortable = lemma.lower()
        lemma_sortable = lemma_sortable.replace('Å', 'å')
        lemma_sortable = lemma_sortable.replace('Ä', 'ä')
        lemma_sortable = lemma_sortable.replace('Ö', 'ö')
        lemma_sortable = lemma_sortable.replace('?', '')
        lemma_sortable = lemma_sortable.replace(',', '')
        lemma_sortable = lemma_sortable.replace('(', '')
        lemma_sortable = lemma_sortable.replace(')', '')
        lemma_sortable = lemma_sortable.replace('-', '')
        return lemma_sortable

    def spolar(self):
        if not hasattr(self, 'spolarna'):
            allSpolar = self.spole_set.all();
            if len(allSpolar) == 0: # Artikeln skapades just, vi fejkar en första spole
                ok = Typ.objects.get(kod = 'OK')
                spole = Spole(typ = ok, text = '', pos = 0)
                self.spolarna = [spole]
            else:
                self.spolarna = allSpolar

        return self.spolarna

    def number_moments(self, type):
        if type == 'M1':
            format = '%d'
            offset = 1
        elif type == 'M2':
            format = '%c'
            offset = 97
        if len(self.moments[type]) > 1:
            for m in range(len(self.moments[type])):
                self.moments[type][m].text = format % (m + offset)
                self.moments[type][m].display = True

    def analyse_spole(self, i):
        spole = self.get_spole(i)
        bits = split(u'¶', spole.text)
        if len(bits) == 1:
            self.append_fjäder(spole)
        else:
            huvudtyp = spole.getType()
            for bit in bits:
                if bits.index(bit) > 0:
                    i += 1
                    if i < len(self.spolar()):
                        self.append_fjäder(Fjäder(self.get_spole(i)), True)
                if bit:
                    fjäder = Fjäder(huvudtyp, bit)
                    if self.preventnextspace and bits.index(bit) == 0:
                        fjäder.preventspace = True
                    self.append_fjäder(fjäder, True)
        return i

    def append_fjäder(self, data, skipmoments = False):
        if skipmoments:
            segment = data
        else:
            segment = Fjäder(data)
            segment.preventspace = self.preventnextspace
            if segment.ism1():
                self.number_moments('M2')
                self.moments['M2'] = []
                self.moments['M1'].append(segment)
            elif segment.ism2():
                self.moments['M2'].append(segment)
        self.fjädrar.append(segment)

    def flush_landskap(self):
        sorted_landskap = sorted(self.landskap, key = Landskap.key)
        if len(sorted_landskap) > 1:
            # Remove province duplicates.
            index = 0;
            while index < (len(sorted_landskap) - 1):
                if sorted_landskap[index].abbrev == sorted_landskap[index + 1].abbrev:
                    del sorted_landskap[index]
                else:
                    index += 1
        sorted_landskap = Landskap.reduce_landskap(sorted_landskap)
        for ls in sorted_landskap:
           fjäder = Fjäder(ArticleTypeManager.get_type_by_abbreviation('g'), ls.abbrev)
           if self.preventnextspace and sorted_landskap.index(ls) == 0:
               fjäder.preventspace = True
           self.append_fjäder(fjäder, True)
        self.landskap = []

    def collect(self):
        self.fjädrar = []
        self.preventnextspace = False
        i = 0
        state = 'ALLMÄNT'
        self.moments = { 'M1': [], 'M2': [] }
        self.landskap = []
        while i < len(self.spolar()):
            spole = self.get_spole(i)
            if spole.isgeo() and spole.text == "":
                i += 1
                continue
            if state == 'ALLMÄNT':
                if spole.isgeo():
                    self.landskap = [Landskap(spole.text)]
                    state = 'GEOGRAFI'
                else:
                    i = self.analyse_spole(i)
                    self.preventnextspace = spole.isleftdelim()
                    # state är fortfarande 'ALLMÄNT'
            elif state == 'GEOGRAFI':
                if spole.isgeo():
                    self.landskap.append(Landskap(spole.text))
                    # state är fortfarande 'GEOGRAFI'
                else:
                    self.flush_landskap()
                    i = self.analyse_spole(i) # För pilcrow i ”hårgård” och ”häringa”
                    self.preventnextspace = spole.isleftdelim()
                    state = 'ALLMÄNT'
            i += 1
        if self.landskap: # För landskapsnamnet på slutet av ”häringa”, efter bugfixet ovan
            self.flush_landskap()
        self.number_moments('M1')
        self.number_moments('M2')
        self.moments = { 'M1': [], 'M2': [] }

    def collect2(self, spolar):
        self.spolarna = spolar
        self.collect()

    def update(self, post_data):
        order = post_data['order'].split(',')
        stickord = post_data['stickord']
        if self.lemma != stickord:
            self.lemma = stickord
            self.lemma_sortable = Artikel.get_lemma_sortable(stickord)
            self.save()
        d = [[post_data['type-' + key].strip(), post_data['value-' + key].strip()] for key in order]
        gtype = Typ.objects.get(kod='g')
        for i in range(len(d)):
            bit = d[i]
            try:
                type = Typ.objects.get(kod=bit[0])
            except ObjectDoesNotExist:  # FIXME Do something useful!
                # TODO >>> Type.objects.create(abbrev = 'OG', name = 'Ogiltig', id = 63)
                type = Typ.objects.get(kod='OG')
            text = bit[1]
            if type == gtype and text.title() in Landskap.short_abbrev.keys():
                text = Landskap.short_abbrev[text.title()]
            data = self.get_spole(i)
            if data and (data.artikel_id != None):
                data.pos = i
                data.typ = type
                data.text = text
                data.save()
            else:
                Spole.objects.create(artikel=self, typ=type, pos=i, text=text)
        self.spole_set.filter(pos__gte=len(d)).delete()


    @staticmethod
    def update_lemma_sortable():
        articles = Artikel.objects.all()
        for article in articles:
            article.lemma_sortable = Artikel.get_lemma_sortable(article.lemma)
            article.save()


class Segment(): # Fjäder!
    def __init__(self, type, text = None):
        self.display = None
        if text != None:
            self.type = type
            self.text = text
        else: # type is actually a Data object
            self.type = type.typ
            self.text = type.text

    def __str__(self):
        return self.type.__str__() + ' ' + self.text

    def isgeo(self):
        return self.type.isgeo()

    def isleftdelim(self):
        if type(self.type) == 'unicode':
            return self.type in ['vh', 'vr']
        else:
            return self.type.isleftdelim()

    def isrightdelim(self):
        if type(self.type) == 'unicode' or str(type(self.type)) == "<type 'unicode'>":
            return self.type in ['hh', 'hr', 'ip', 'ko'] # KO is convenient
        else:
            return self.type.isrightdelim()

    def ismoment(self):
        return self.ism1() or self.ism2()

    def ism1(self):
        return self.type.ism1()

    def ism2(self):
        return self.type.ism2()

    def output(self, outfile, next):
        outfile.write(('\SDL:%s{' % self.type).encode('UTF-8'))
        outfile.write(self.text.replace(u'\\', '\\backslash ').encode('UTF-8'))
        outfile.write('}')
        if not next.isrightdelim():
            outfile.write(' ')

    # TODO Method hyphornot()

    def format(self):
        if self.type.__str__().upper() == 'VH':
            return '['
        elif self.type.__str__().upper() == 'HH':
            return ']'
        elif self.type.__str__().upper() == 'VR':
            return '('
        elif self.type.__str__().upper() == 'HR':
            return ')'
        elif self.type.__str__().upper() == 'ÄV':
            return 'äv.'
        else:
            return self.text.strip()

class Typ(models.Model):
    kod = models.CharField(max_length = 5)
    namn = models.CharField(max_length = 30)
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)

    def __str__(self):
        return self.kod.upper()

    def isgeo(self):
        return self.kod == 'g'

    def format(self):
        out = self.__str__()
        out += (4 - len(out)) * '\xa0'
        return out

class ArticleSearchCriteria:
    def __init__(self):
        self.compare_type = "StartsWith"
        self.search_string = ""
        self.search_type = "Lemma"

class ArticleManager:

    @staticmethod
    def get_article_item_types(search_criteria_list):
        articleItemTypeIds = []
        has_article_item_search_criteria = False
        has_article_item_type_search_criteria = False
        for search_criteria in search_criteria_list:
            if ((search_criteria.search_type == 'ArticleItemType') and
                (not (is_empty_string(search_criteria.search_string)))):
                has_article_item_type_search_criteria = True
            if (search_criteria.search_type == 'ArticleItem'):
                has_article_item_search_criteria = True

        if (has_article_item_search_criteria and
            has_article_item_type_search_criteria):
            for search_criteria in search_criteria_list:
                if ((search_criteria.search_type == 'ArticleItemType') and
                    (not (is_empty_string(search_criteria.search_string)))):
                    if (search_criteria.compare_type == 'Contains'):
                        articleItemTypes = Typ.objects.filter(kod__icontains=search_criteria.search_string)
                    elif (search_criteria.compare_type == 'EndWith'):
                        articleItemTypes = Typ.objects.filter(kod__iendswith=search_criteria.search_string)
                    elif (search_criteria.compare_type == 'EqualTo'):
                        articleItemTypes = Typ.objects.filter(kod__iexact=search_criteria.search_string)
                    elif (search_criteria.compare_type == 'StartsWith'):
                        articleItemTypes = Typ.objects.filter(kod__istartswith=search_criteria.search_string)
                    for articleItemType in articleItemTypes:
                        articleItemTypeIds.append(articleItemType.id)
        else:
            articleItemTypeIds = None
        return articleItemTypeIds

    @staticmethod
    def get_articles_by_search_criteria(search_criteria_list):
        if (ArticleManager.is_empty_search_criteria(search_criteria_list)):
            return []

        article_search_result = None
        article_item_types = ArticleManager.get_article_item_types(search_criteria_list)
        for search_criteria in search_criteria_list:
            if ((not (is_empty_string(search_criteria.search_string))) and
                (not ((search_criteria.search_type == 'ArticleItemType') and
                      (not (article_item_types is None))))):
                articles = ArticleManager.get_articles_by_one_search_criteria(search_criteria, article_item_types)
                if (article_search_result is None):
                    article_search_result = articles
                else:
                    article_search_result = ArticleManager.get_subset(article_search_result, articles)

        if (article_search_result is None):
            article_search_result = []
        return article_search_result

    @staticmethod
    def get_articles_by_one_search_criteria(search_criteria, article_item_types):
        articles = []
        search_articles = None
        if (not (is_empty_string(search_criteria.search_string))):
            if ((search_criteria.search_type == 'Lemma') or
                (search_criteria.search_type == 'All')):
                if (search_criteria.compare_type == 'Contains'):
                    search_articles = Artikel.objects.filter(lemma_sortable__icontains = search_criteria.search_string)
                elif (search_criteria.compare_type == 'EndWith'):
                    search_articles = Artikel.objects.filter(lemma_sortable__iendswith=search_criteria.search_string)
                elif (search_criteria.compare_type == 'EqualTo'):
                    search_articles = Artikel.objects.filter(lemma_sortable__iexact=search_criteria.search_string)
                elif (search_criteria.compare_type == 'StartsWith'):
                    search_articles = Artikel.objects.filter(lemma_sortable__istartswith=search_criteria.search_string)
                articles += search_articles

            if ((search_criteria.search_type == 'ArticleItem') or
                (search_criteria.search_type == 'All')):
                if (search_criteria.compare_type == 'Contains'):
                    if (article_item_types is None):
                        articleItems = Spole.objects.filter(text__icontains=search_criteria.search_string).select_related('artikel')
                    else:
                        articleItems = Spole.objects.filter(text__icontains=search_criteria.search_string, typ_id__in=article_item_types).select_related('artikel')
                elif (search_criteria.compare_type == 'EndWith'):
                    if (article_item_types is None):
                        articleItems = Spole.objects.filter(text__iendswith=search_criteria.search_string).select_related('artikel')
                    else:
                        articleItems = Spole.objects.filter(text__iendswith=search_criteria.search_string, typ_id__in=article_item_types).select_related('artikel')
                elif (search_criteria.compare_type == 'EqualTo'):
                    if (article_item_types is None):
                        articleItems = Spole.objects.filter(text__iexact=search_criteria.search_string).select_related('artikel')
                    else:
                        articleItems = Spole.objects.filter(text__iexact=search_criteria.search_string, typ_id__in=article_item_types).select_related('artikel')
                elif (search_criteria.compare_type == 'StartsWith'):
                    if (article_item_types is None):
                        articleItems = Spole.objects.filter(text__istartswith=search_criteria.search_string).select_related('artikel')
                    else:
                        articleItems = Spole.objects.filter(text__istartswith=search_criteria.search_string, typ_id__in=article_item_types).select_related('artikel')

                search_articles = [articleItem.artikel for articleItem in articleItems]
                articles += search_articles

            if (((search_criteria.search_type == 'ArticleItemType') or
                 (search_criteria.search_type == 'All')) and
                (article_item_types is None)):
                if (search_criteria.compare_type == 'Contains'):
                    articleItemTypes = Typ.objects.filter(kod__icontains=search_criteria.search_string)
                elif (search_criteria.compare_type == 'EndWith'):
                    articleItemTypes = Typ.objects.filter(kod__iendswith=search_criteria.search_string)
                elif (search_criteria.compare_type == 'EqualTo'):
                    articleItemTypes = Typ.objects.filter(kod__iexact=search_criteria.search_string)
                elif (search_criteria.compare_type == 'StartsWith'):
                    articleItemTypes = Typ.objects.filter(kod__istartswith=search_criteria.search_string)
                articleItemTypeIds = []
                for articleItemType in articleItemTypes:
                    articleItemTypeIds.append(articleItemType.id)
                articleItems = Spole.objects.filter(typ_id__in=articleItemTypeIds).select_related('artikel')
                search_articles = [articleItem.artikel for articleItem in articleItems]
                articles += search_articles

        if (len(articles) > 1):
            # Get distinct articles.
            articlesDictionary = {}
            for article in articles:
                if (not (article.id in articlesDictionary)):
                    articlesDictionary[article.id] = article

            # Sorted articles from database.
            search_articles = Artikel.objects.filter(id__in=articlesDictionary.keys())
            articles = []
            articles += search_articles

        return articles


    @staticmethod
    def get_subset(articles1, articles2):
        if (len(articles1) < 1) or (articles1[0] is None):
            return []
        if (len(articles2) < 1) or (articles2[0] is None):
            return []

        article_dictionary = {}
        for article in articles2:
            article_dictionary[article.id] = article

        index = len(articles1) - 1
        while (index >= 0):
            if (not (articles1[index].id in article_dictionary)):
                articles1.remove(articles1[index])
            index -= 1

        return articles1

    @staticmethod
    def is_empty_search_criteria(search_criteria_list):
        if (search_criteria_list is None):
            return True;
        if (len(search_criteria_list) < 1):
            return True;
        is_search_string_found = False
        for search_criteria in search_criteria_list:
            if (is_not_empty_string(search_criteria.search_string)):
                is_search_string_found = True
        return not is_search_string_found


article_types_by_abbreviation = {}
article_types_by_id = {}

class ArticleTypeManager:

    @staticmethod
    def get_type_by_abbreviation(abbreviation):
        ArticleTypeManager.init_types()

        # Get requested article type.
        if abbreviation in article_types_by_abbreviation:
            return article_types_by_abbreviation[abbreviation]
        else:
            raise ObjectDoesNotExist("No article type with abbreviation = " + abbreviation)

    @staticmethod
    def get_type_by_id(id):
        ArticleTypeManager.init_types()

        # Get requested article type.
        if id in article_types_by_id:
            return article_types_by_id[id]
        else:
            raise ObjectDoesNotExist("No article type with id = " + str(id))

    @staticmethod
    def init_types():
        if len(article_types_by_id) == 0:
            # Get all article types.
            allArticleTypes = Typ.objects.all()
            for articleType in allArticleTypes:
                article_types_by_abbreviation[articleType.kod] = articleType
                article_types_by_id[articleType.id] = articleType

class Spole(models.Model):
    text = models.CharField(max_length = 2000)
    pos = models.SmallIntegerField()
    artikel = models.ForeignKey(Artikel)
    typ = models.ForeignKey(Typ)
    skapat = models.DateTimeField(auto_now_add = True)
    uppdaterat = models.DateTimeField(auto_now = True)

    class Meta:
        ordering = ('pos',)

    def __str__(self):
        return self.getType().__str__() + ' ' + self.text

    def webstyle(self):
        self.webstyles[self.getType().__str__()]

    def printstyle(self):
        self.printstyles[self.getType().__str__()]

    def isgeo(self):
        return self.getType().isgeo()
        # return self.typ.isgeo()

    def isleftdelim(self):
        return self.getType().kod.upper() in ['VR', 'VH']

    def getType(self):
        return ArticleTypeManager.get_type_by_id(self.typ_id)


class Fjäder:
    def __init__(self, spole, text = None):
        self.display = None # FIXME Måste inte ändras för KO
        if text: # spole är egentligen en Typ
            self.typ = spole.kod.upper()
            self.text = text.strip()
        else:
            self.typ = spole.getType().kod.upper()
            self.text = spole.text.strip()
        self.preventspace = False

    def isrightdelim(self):
        return self.typ in ['HH', 'HR', 'IP', 'KO'] or match(r'^,', self.text)

    def ismoment(self):
        return self.ism1() or self.ism2()

    def ism1(self):
        return self.typ == 'M1'

    def ism2(self):
        return self.typ == 'M2'

    def format(self): # FIXME Också för P:er från ÖVP?
        if self.typ == 'VR':
            return '('
        elif self.typ == 'HR':
            return ')'
        elif self.typ == 'VH':
            return '['
        elif self.typ == 'HH':
            return ']'
        elif self.typ == 'ÄV':
            return 'äv.'
        else:
            return self.text

    def type(self): # FIXME Remove later!
        return self.typ

    def setspace(self):
        return not self.preventspace and not self.isrightdelim() # TODO Lägga till det med den föregående fjädern

class Landskap():
    ordning = [
        u'Skåne', u'Blek', u'Öland', u'Smål', u'Hall', u'Västg', u'Boh', u'Dalsl', u'Gotl', u'Östg', # 0-9
        u'Götal', # 10
        u'Sörml', u'Närke', u'Värml', u'Uppl', u'Västm', u'Dal', # 11 - 16
        u'Sveal', # 17
        u'Gästr', u'Häls', u'Härj', u'Med', u'Jämtl', u'Ång', u'Västb', u'Lappl', u'Norrb', # 18 - 26
        u'Norrl', # 27
    ]

    # TODO: Something more object-oriented (methods short, long, rank)
    short_abbrev = {
      u'Sk' : u'Skåne',
      u'Bl' : u'Blek',
      u'Öl' : u'Öland',
      u'Sm' : u'Smål',
      u'Ha' : u'Hall',
      u'Vg' : u'Västg',
      u'Bo' : u'Boh',
      u'Dsl': u'Dalsl',
      u'Gl' : u'Gotl',
      u'Ög' : u'Östg',
      u'Götal' : u'Götal',
      u'Sdm': u'Sörml',
      u'Nk' : u'Närke',
      u'Vrm': u'Värml',
      u'Ul' : u'Uppl',
      u'Vstm' : u'Västm',
      u'Dal': u'Dal',
      u'Sveal' : u'Sveal',
      u'Gst': u'Gästr',
      u'Hsl': u'Häls',
      u'Hrj': u'Härj',
      u'Mp' : u'Med',
      u'Jl' : u'Jämtl',
      u'Åm' : u'Ång',
      u'Vb' : u'Västb',
      u'Lpl': u'Lappl',
      u'Nb' : u'Norrb',
      u'Norrl' : u'Norrl',
    }

    def cmp(self, other):
        if self.abbrev in self.ordning and other.abbrev in self.ordning:
            return self.cmp(self.ordning.index(self.abbrev), self.ordning.index(other.abbrev))
        else:
            return 0

    @staticmethod
    def key(self):
        if self.abbrev in self.ordning:
            return self.ordning.index(self.abbrev)
        else:
            return -1 # Så det är lättare att se dem

    def __init__(self, abbrev):
        self.abbrev = abbrev.capitalize()

    def __str__(self):
        return self.abbrev

    @staticmethod
    def reduce_landskap(lista):
        antal_per_del = { 'Götal' : 7, 'Sveal' : 4, 'Norrl' : 6 }
        delar = ['Götal', 'Sveal', 'Norrl']
        landskap_per_del = { 'Götal' : [], 'Sveal' : [], 'Norrl' : [] }
        for landskap in lista:
            if Landskap.ordning.index(landskap.abbrev) < Landskap.ordning.index('Götal'):
                landskap_per_del['Götal'].append(landskap)
            elif Landskap.ordning.index(landskap.abbrev) < Landskap.ordning.index('Sveal'):
                landskap_per_del['Sveal'].append(landskap)
            else:
                landskap_per_del['Norrl'].append(landskap)
            # print(landskap)
        # print(len(landskap_per_del['Götal']))
        utskriftsform = []
        for landsdel, landskapen in landskap_per_del.items():
            if len(landskapen) >= antal_per_del[landsdel]:
                utskriftsform.append(Landskap(landsdel))
            else:
                utskriftsform += sorted(landskapen, key = Landskap.key)

        return utskriftsform

class Exporter:
    def __init__(self, format):
        self.format = format
        initialisers = {
            'pdf' : self.start_pdf,
            'odt' : self.start_odf,
            'docx' : self.start_docx,
        }

        generators = {
            'pdf' : self.generate_pdf_paragraph,
            'odt' : self.generate_odf_paragraph,
            'docx' : self.generate_docx_paragraph,
        }

        savers = {
            'pdf' : self.save_pdf,
            'odt' : self.save_odf,
            'docx' : self.save_docx,
        }

        self.start_document = initialisers[format]
        self.generate_paragraph = generators[format]
        self.save_document = savers[format]

    def export_letter(self, letter):
        # Get data from database.
        artiklar = []
        hel_bokstav = Artikel.objects.filter(lemma__startswith=letter)
        artiklar += hel_bokstav

        if len(artiklar) >= 1:
            artikelIds = []
            for artikel in artiklar:
                artikelIds.append(artikel.id)
            spoles = []
            spolar = Spole.objects.filter(artikel_id__in=artikelIds)
            spoles += spolar
            articleSpoles = {}
            for id in artikelIds:
                articleSpoles[id] = []
            for spole in spoles:
                articleSpoles[spole.artikel_id].append(spole)

            # Create file with article information.
            self.dirname = mkdtemp('', 'SDLartikel')
            self.filename = join(self.dirname, 'sdl.%s' % self.format)
            self.start_document()
            filename = letter + datetime.datetime.now().strftime("%Y%m%d") + '.' + self.format
            for artikel in artiklar:
                artikel.collect2(articleSpoles[artikel.id])
                self.generate_paragraph(artikel)
            self.save_document()
            staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
            rename(self.filename, join(staticpath, filename))

    # Create files, one file for each letter in the swedish alphabet.
    def export_letters(self):
        letter = 'a'
        while letter <= 'z':
            self.export_letter(letter)
            letter = chr(ord(letter) + 1)
        letter = 'å'
        self.export_letter(letter)
        letter = 'ä'
        self.export_letter(letter)
        letter = 'ö'
        self.export_letter(letter)

    def generate_pdf_paragraph(self, artikel):
        self.document.write("\\hskip-0.5em")
        if artikel.rang > 0:
            self.document.write('\lohi[left]{}{\SDL:SO{%d}}' % artikel.rang) # TODO Real superscript
        self.document.write("\\SDL:SO{%s}" % artikel.lemma)
        for segment in artikel.fjädrar: # TODO Handle moments!  segment.ismoment and segment.display
            if not segment.preventspace and not segment.isrightdelim():
                self.document.write(' ') # FIXME But not if previous segment is left delim!
            type = segment.typ
            text = segment.format().replace(u'\\', '\\textbackslash ').replace('~', '{\\char"7E}')
            self.document.write(('\\SDL:%s{%s}' % (type, text)))
        self.document.write("\\par")

    def start_pdf(self):
        self.document = open(self.filename.replace('.pdf', '.tex'), 'w')
        self.document.write("\\mainlanguage[sv]")
        self.document.write("\\setupbodyfont[pagella, 12pt]\n")
        self.document.write("\\setuppagenumbering[state=stop]\n")
        with open(join(dirname(abspath(__file__)), 'lib', 'sdl-setup.tex'), 'r', encoding = 'UTF-8') as file:
            self.document.write(file.read())

        self.document.write("""
            \\starttext
            """)

        self.document.write("\\startcolumns[n=2,balance=yes]\n")

    def save_pdf(self):
        self.document.write("\\stopcolumns\n")
        self.document.write("\\stoptext\n")
        self.document.close()

        chdir(self.dirname)
        environ['PATH'] = "%s:%s" % (settings.CONTEXT_PATH, environ['PATH'])
        environ['TMPDIR'] = '/tmp'
        popen("context --batchmode sdl.tex").read()

    def add_odf_style(self, type, xmlchunk):
        self.document.inject_style("""
            <style:style style:name="%s" style:family="text">
                <style:text-properties %s />
            </style:style>
            """ % (type, xmlchunk))

    def start_odf(self):
        self.document = ezodf.newdoc(doctype = 'odt', filename = self.filename)
        self.add_odf_style('SO', 'fo:font-weight="bold"')
        self.add_odf_style('OK', 'fo:font-size="9pt"')
        self.add_odf_style('G', 'fo:font-size="9pt"')
        self.add_odf_style('DSP', 'fo:font-style="italic"')
        self.add_odf_style('TIP', 'fo:font-size="9pt"')
        # IP
        self.add_odf_style('M1', 'fo:font-weight="bold"')
        self.add_odf_style('M2', 'fo:font-weight="bold"')
        # VH, HH, VR, HR
        self.add_odf_style('REF', 'fo:font-weight="bold"')
        self.add_odf_style('FO', 'fo:font-style="italic"')
        self.add_odf_style('TIK', 'fo:font-style="italic" fo:font-size="9pt"')
        self.add_odf_style('FLV', 'fo:font-weight="bold" fo:font-size="9pt"')
        # ÖVP. Se nedan!
        # BE, ÖV
        # ÄV, ÄVK se nedan
        # FOT
        self.add_odf_style('GT', 'fo:font-size="9pt"')
        self.add_odf_style('SOV', 'fo:font-weight="bold"')
        # TI
        # HV, INT
        self.add_odf_style('OKT', 'fo:font-size="9pt"')
        # VS
        self.add_odf_style('GÖ', 'fo:font-size="9pt"')
        # GP, UST
        self.add_odf_style('US', 'fo:font-style="italic"')
        # GÖP, GTP, NYR, VB
        self.add_odf_style('OG', 'style:text-line-through-style="solid"')
        self.add_odf_style('SP', 'fo:font-style="italic"')

    def save_odf(self):
        self.document.save()

    def generate_odf_paragraph(self, artikel):
        paragraph = ezodf.Paragraph()
        paragraph += ezodf.Span(artikel.lemma, style_name = 'SO') # TODO Homografnumrering!
        for segment in artikel.fjädrar:
            type = segment.typ
            if not type == 'KO':
                if not segment.preventspace and not segment.isrightdelim():
                    paragraph += ezodf.Span(' ')
                paragraph += ezodf.Span(segment.format(), style_name = type)
        self.document.body += paragraph

    def start_docx(self):
        self.document = docx.Document()
        self.add_docx_styles()
        return self.document

    def save_docx(self):
        self.document.save(self.filename)

    def generate_docx_paragraph(self, artikel):
        paragraph = self.document.add_paragraph()
        if artikel.rang > 0:
            paragraph.add_run(str(artikel.rang), style = 'SO').font.superscript = True
        paragraph.add_run(artikel.lemma, style = 'SO')
        for segment in artikel.fjädrar:
            type = segment.typ
            if not type == 'KO':
                if not segment.preventspace and not segment.isrightdelim():
                    paragraph.add_run(' ', style = self.document.styles[type])
                if type == 'ÖVP': # TODO Något bättre
                    run = '(' + segment.format() + ')'
                else:
                    run = segment.format()
                paragraph.add_run(run, style = self.document.styles[type])

    def add_docx_styles(self): # TODO Keyword arguments?
        self.add_docx_style('SO', False, True, 12)
        self.add_docx_style('OK', False, False, 9)
        self.add_docx_style('G', False, False, 9)
        self.add_docx_style('DSP', True, False, 12)
        self.add_docx_style('TIP', False, False, 9)
        self.add_docx_style('IP', False, False, 12)
        self.add_docx_style('M1', False, True, 12)
        self.add_docx_style('M2', False, True, 12)
        self.add_docx_style('VH', False, False, 12)
        self.add_docx_style('HH', False, False, 12)
        self.add_docx_style('VR', False, False, 12)
        self.add_docx_style('HR', False, False, 12)
        self.add_docx_style('REF', False, True, 12)
        self.add_docx_style('FO', True, False, 12)
        self.add_docx_style('TIK', True, False, 9)
        self.add_docx_style('FLV', False, True, 9)
        self.add_docx_style('ÖVP', False, False, 12)
        self.add_docx_style('BE', False, False, 12)
        self.add_docx_style('ÖV', False, False, 12)
        self.add_docx_style('ÄV', False, False, 12) # FIXME Skriv “även” :-)
        self.add_docx_style('ÄVK', True, False, 12) # FIXME Skriv även även här ;-)
        self.add_docx_style('FOT', True, False, 12)
        self.add_docx_style('GT', False, False, 9)
        self.add_docx_style('SOV', False, True, 12)
        for style in ('TI', 'HV', 'INT'):
            self.add_docx_style(style)
        self.add_docx_style('OKT', False, False, 9)
        self.add_docx_style('VS')
        self.add_docx_style('GÖ', False, False, 9)
        self.add_docx_style('GP')
        self.add_docx_style('UST', True, False, 12)
        self.add_docx_style('US', True, False, 12)
        for style in ('GÖP', 'GTP', 'NYR', 'VB'):
            self.add_docx_style(style)
        OG = self.document.styles.add_style('OG', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        OG.font.strike = True
        self.add_docx_style('SP', True, False, 12)
        self.add_docx_style('M0', False, True, 18)
        self.add_docx_style('M3', True, False, 6)

    def add_docx_style(self, type, italic = False, bold = False, size = 12):
        style = self.document.styles.add_style(type, docx.enum.style.WD_STYLE_TYPE.CHARACTER)
        style.base_style = self.document.styles['Default Paragraph Font']
        if italic:
            style.font.italic = True
        if bold:
            style.font.bold = True
        style.font.size = docx.shared.Pt(size)

    def export(self, ids):
        self.dirname = mkdtemp('', 'SDLartikel')
        self.filename = join(self.dirname, 'sdl.%s' % self.format)
        self.start_document()
        if (len(ids) == 0) or (len(ids) == 1 and ids[0] == ''):
            # Nothing to do. No articles should be output to file.
            pass
        elif len(ids) == 1:
            artikel = Artikel.objects.get(id = ids[0])
            artikel.collect()
            self.generate_paragraph(artikel)
            filename = '%s-%s.%s' % (ids[0], artikel.lemma, self.format)
        elif len(ids) > 1:
            # Get data from database.
            artiklar = []
            hel_bokstav = Artikel.objects.filter(id__in=ids)
            artiklar += hel_bokstav
            spoles = []
            # spolar = Spole.objects.filter(artikel_id__in=ids).order_by('pos').asc()
            spolar = Spole.objects.filter(artikel_id__in=ids)
            spoles += spolar
            articleSpoles = {}
            for id in ids:
                articleSpoles[int(id)] = []
            for spole in spoles:
                articleSpoles[spole.artikel_id].append(spole)

            # Write article information to file.
            filename = 'sdl-utdrag.%s' % self.format # FIXME Unikt namn osv.
            for artikel in artiklar:
                artikel.collect2(articleSpoles[artikel.id])
                self.generate_paragraph(artikel)
        self.save_document()
        staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
        copyfile(self.filename, join(staticpath, filename))

        return join('ord', filename)

    def export_articles(self, articles, userName):
        # Get temp file.
        self.dirname = mkdtemp('', 'SDLartikel')
        self.filename = join(self.dirname, 'sdl.%s' % self.format)
        self.start_document()

        # Write article information to file.
        filename = f'{userName}-sdl.{self.format}'
        for article in articles:
            self.generate_paragraph(article)
        self.save_document()

        # Copy temp file to output file.
        staticpath = join(dirname(abspath(__file__)), 'static', 'ord')
        filepath = join(staticpath, filename)
        copyfile(self.filename, filepath)

        return filepath

